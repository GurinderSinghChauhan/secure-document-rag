from dataclasses import dataclass, field
from io import BytesIO
import warnings

from docx import Document as DocxDocument
from fastapi import HTTPException, status
from PIL import Image


@dataclass(frozen=True)
class VisualAsset:
    content: bytes
    media_type: str
    location: str
    description_indexed: bool = False


@dataclass
class ParsedDocument:
    text: str
    visuals: list[VisualAsset] = field(default_factory=list)
    table_count: int = 0
    described_visual_count: int = 0


def _cell_text(value: object) -> str:
    return " ".join(str(value or "").replace("|", "\\|").split())


def table_to_markdown(rows: list[list[object]], label: str) -> str:
    normalized = [[_cell_text(cell) for cell in row[:20]] for row in rows]
    normalized = [row for row in normalized if any(row)]
    if not normalized:
        return ""
    column_count = max(len(row) for row in normalized)
    padded = [row + [""] * (column_count - len(row)) for row in normalized]
    header = padded[0]
    separator = ["---"] * column_count
    markdown_rows = [header, separator, *padded[1:]]
    lines = [f"| {' | '.join(row)} |" for row in markdown_rows]
    return f"[{label}]\n" + "\n".join(lines)


def _image_to_visual(image: Image.Image, location: str) -> VisualAsset:
    image.seek(0)
    image.thumbnail((1600, 1600))
    normalized = image.convert("RGBA" if image.mode in {"RGBA", "LA"} else "RGB")
    output = BytesIO()
    normalized.save(output, format="PNG", optimize=True)
    return VisualAsset(content=output.getvalue(), media_type="image/png", location=location)


def normalize_visual(content: bytes, location: str, description_indexed: bool = False) -> VisualAsset | None:
    try:
        with warnings.catch_warnings():
            warnings.simplefilter("error", Image.DecompressionBombWarning)
            with Image.open(BytesIO(content)) as image:
                image.load()
                visual = _image_to_visual(image, location)
                return VisualAsset(visual.content, visual.media_type, visual.location, description_indexed)
    except Exception:
        return None


def _extract_docx(content: bytes, max_visuals: int) -> ParsedDocument:
    document = DocxDocument(BytesIO(content))
    sections = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_count = 0
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        markdown = table_to_markdown(rows, f"Table {table_index}")
        if markdown:
            sections.append(markdown)
            table_count += 1

    visuals: list[VisualAsset] = []
    seen_parts: set[str] = set()
    for relationship in document.part.rels.values():
        if len(visuals) >= max_visuals or not relationship.reltype.endswith("/image"):
            continue
        part_name = str(relationship.target_part.partname)
        if part_name in seen_parts:
            continue
        seen_parts.add(part_name)
        visual = normalize_visual(relationship.target_part.blob, f"DOCX embedded image {len(visuals) + 1}")
        if visual:
            visuals.append(visual)
    return ParsedDocument(text="\n\n".join(sections), visuals=visuals, table_count=table_count)


def extract_document(content: bytes, content_type: str, max_visuals: int = 40) -> ParsedDocument:
    try:
        if content_type.startswith("text/plain"):
            return ParsedDocument(text=content.decode("utf-8"))
        if content_type.startswith("application/vnd.openxmlformats-officedocument.wordprocessingml.document"):
            return _extract_docx(content, max_visuals)
        if content_type.startswith("image/"):
            visual = normalize_visual(content, "Uploaded image")
            if visual:
                return ParsedDocument(text="", visuals=[visual])
            raise ValueError("Unsupported image format")
    except HTTPException:
        raise
    except Exception as error:
        raise HTTPException(status_code=status.HTTP_422_UNPROCESSABLE_ENTITY, detail="Unable to parse document") from error
    raise HTTPException(
        status_code=status.HTTP_415_UNSUPPORTED_MEDIA_TYPE,
        detail="This document type requires the configured MinerU parser",
    )


def extract_text(content: bytes, content_type: str) -> str:
    return extract_document(content, content_type).text

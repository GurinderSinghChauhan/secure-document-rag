import base64
from io import BytesIO

from docx import Document
from docx.shared import Inches
from pypdf import PdfWriter

from app.document_parser import extract_document, table_to_markdown


PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mNk+A8AAQUBAScY42YAAAAASUVORK5CYII="
)


def test_table_to_markdown_preserves_cells_and_escapes_pipes():
    markdown = table_to_markdown([["Name", "Value"], ["Risk | level", 5]], "Table 1")

    assert "[Table 1]" in markdown
    assert "Risk \\| level" in markdown
    assert "| 5 |" in markdown


def test_extract_docx_finds_paragraph_table_and_image():
    document = Document()
    document.add_paragraph("Clinical review")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Finding"
    table.cell(0, 1).text = "Status"
    table.cell(1, 0).text = "Consent"
    table.cell(1, 1).text = "Complete"
    document.add_picture(BytesIO(PNG_BYTES), width=Inches(0.2))
    output = BytesIO()
    document.save(output)

    parsed = extract_document(
        output.getvalue(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert "Clinical review" in parsed.text
    assert "Consent" in parsed.text
    assert parsed.table_count == 1
    assert len(parsed.visuals) == 1


def test_extract_standalone_image_as_visual():
    parsed = extract_document(PNG_BYTES, "image/png")

    assert parsed.text == ""
    assert len(parsed.visuals) == 1
    assert parsed.visuals[0].location == "Uploaded image"


def test_extract_pdf_fallback_when_mineru_is_disabled():
    output = BytesIO()
    writer = PdfWriter()
    writer.add_blank_page(width=72, height=72)
    writer.write(output)

    parsed = extract_document(output.getvalue(), "application/pdf")

    assert parsed.text == ""
    assert parsed.visuals == []
